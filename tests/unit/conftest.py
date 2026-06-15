import io

import pytest
from docx import Document as DocxDocument


@pytest.fixture
def docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("Hello DOCX")
    doc.add_paragraph("Second paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(text: str) -> bytes:
    # Minimal valid PDF with a single page and extractable text.
    content_stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode("latin-1")

    objects: list[bytes] = [
        b"<</Type /Catalog /Pages 2 0 R>>",
        b"<</Type /Pages /Kids [3 0 R] /Count 1>>",
        b"<</Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
        b" /Contents 4 0 R /Resources<</Font<</F1 5 0 R>>>>>>",
        b"<</Length "
        + str(len(content_stream)).encode()
        + b">>\nstream\n"
        + content_stream
        + b"\nendstream",
        b"<</Type /Font /Subtype /Type1 /BaseFont /Helvetica>>",
    ]

    header = b"%PDF-1.4\n"
    body = header
    offsets: list[int] = []

    for i, obj_body in enumerate(objects, start=1):
        offsets.append(len(body))
        body += f"{i} 0 obj\n".encode() + obj_body + b"\nendobj\n"

    xref_pos = len(body)
    xref = b"xref\n0 6\n0000000000 65535 f \n"
    for offset in offsets:
        xref += f"{offset:010d} 00000 n \n".encode()

    trailer = (
        b"trailer\n<</Size 6 /Root 1 0 R>>\nstartxref\n" + str(xref_pos).encode() + b"\n%%EOF\n"
    )

    return body + xref + trailer


@pytest.fixture
def pdf_bytes() -> bytes:
    return _build_pdf("Hello PDF")
